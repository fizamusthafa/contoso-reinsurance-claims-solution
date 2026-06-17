import os
from docx import Document

BUILD = r"c:\Users\fizamusthafa\customerdemos\contoso reinsurance claims case\build\botcomponents"

def save(rel, builder):
    path = os.path.join(BUILD, rel)
    doc = Document()
    builder(doc)
    doc.save(path)
    print("wrote", path)

# ---------- Region Mapping (Germany) ----------
def region(doc):
    doc.add_heading("Region Mapping \u2013 Germany", level=1)
    doc.add_paragraph("Contoso Reinsurance covers losses across the German federal states (Bundesl\u00e4nder). "
                      "Detect the region from the loss-location/insured address and the subject line.")
    states = [
        ("Bayern (BY)", "M\u00fcnchen, Munich, N\u00fcrnberg, Augsburg, Regensburg, BY, Bayern, Bavaria"),
        ("Nordrhein-Westfalen (NW)", "K\u00f6ln, Cologne, D\u00fcsseldorf, Dortmund, Essen, Duisburg, NRW, NW"),
        ("Baden-W\u00fcrttemberg (BW)", "Stuttgart, Karlsruhe, Mannheim, Freiburg, Heidelberg, BW"),
        ("Hessen (HE)", "Frankfurt, Wiesbaden, Darmstadt, Kassel, HE"),
        ("Berlin (BE)", "Berlin, BE"),
        ("Hamburg (HH)", "Hamburg, HH"),
        ("Niedersachsen (NI)", "Hannover, Braunschweig, Osnabr\u00fcck, NI"),
        ("Sachsen (SN)", "Dresden, Leipzig, Chemnitz, SN"),
    ]
    for name, ids in states:
        doc.add_heading(name, level=2)
        doc.add_paragraph("Identifiers: " + ids)
        doc.add_paragraph("Primary Language: German (DE)")
    doc.add_paragraph("Note: Detect the language from the content of the claim, not from the region. "
                      "If the loss notification is written in English, set language EN even for a German state. "
                      "If it is predominantly German, set language DE.")

# ---------- Claims Adjuster Assignment Table ----------
def adjusters(doc):
    doc.add_heading("Claims Adjuster Assignment Table", level=1)
    doc.add_paragraph("Contoso Reinsurance claims adjusters. Prefer an adjuster whose region (Bundesland) "
                      "and language match the claim, then balance by caseload and Outlook availability.")
    rows = [
        ("Mario Rogers",   "MarioR@M365CPI77548680.OnMicrosoft.com",  "Senior", "Bayern",              "DE, EN", "8 / 15"),
        ("Lisa Taylor",    "LisaT@M365CPI77548680.OnMicrosoft.com",   "Mid",    "Nordrhein-Westfalen", "DE, EN", "6 / 18"),
        ("Sydney Mattos",  "SydneyM@M365CPI77548680.OnMicrosoft.com", "Junior", "Hessen",              "DE, EN", "4 / 20"),
        ("Daichi Maruyama","DaichiM@M365CPI77548680.OnMicrosoft.com", "Senior", "Baden-W\u00fcrttemberg",   "DE, EN", "11 / 15"),
        ("Cora Thomas",    "CoraT@M365CPI77548680.OnMicrosoft.com",   "Mid",    "Berlin",              "DE, EN", "9 / 18"),
    ]
    for name, email, sen, reg, langs, load in rows:
        doc.add_heading(name, level=2)
        doc.add_paragraph("Email: " + email)
        doc.add_paragraph("Seniority: " + sen)
        doc.add_paragraph("Region: " + reg)
        doc.add_paragraph("Languages: " + langs)
        doc.add_paragraph("Caseload: " + load)
        doc.add_paragraph("Status: Active")
    doc.add_heading("Fallback Adjuster", level=2)
    doc.add_paragraph("If no adjuster is eligible, assign to Amber Rodriguez, "
                      "AmberR@M365CPI77548680.OnMicrosoft.com (Senior, Hamburg, DE/EN).")

# ---------- Cause of Loss Reference ----------
def causes(doc):
    doc.add_heading("Cause of Loss Reference", level=1)
    doc.add_paragraph("Classify the cause of loss / peril from the loss description. "
                      "Return a code, title, and severity. German terms are shown for convenience.")
    groups = [
        ("Natural Catastrophe (High Severity)", [
            "Flood / \u00dcberschwemmung \u2192 NATCAT-FLD \u2014 Flood",
            "Windstorm & Hail / Sturm & Hagel \u2192 NATCAT-WND \u2014 Windstorm & Hail",
            "Earthquake / Erdbeben \u2192 NATCAT-EQ \u2014 Earthquake",
        ]),
        ("Fire (High Severity)", [
            "Fire / Brand \u2192 FIRE-100 \u2014 Fire & Explosion",
            "Explosion / Explosion \u2192 FIRE-110 \u2014 Explosion",
        ]),
        ("Water Damage (Medium Severity)", [
            "Burst pipe / Leitungswasser \u2192 WATER-200 \u2014 Escape of Water",
        ]),
        ("Property Damage (Medium Severity)", [
            "Storm/hail building damage \u2192 PROP-300 \u2014 Property Damage",
            "Impact / collision / Anprall \u2192 PROP-310 \u2014 Impact Damage",
        ]),
        ("Theft & Burglary (Medium Severity)", [
            "Theft / Einbruchdiebstahl \u2192 THEFT-400 \u2014 Theft & Burglary",
        ]),
        ("Business Interruption (High Severity)", [
            "Loss of profits / Betriebsunterbrechung \u2192 BI-500 \u2014 Business Interruption",
        ]),
        ("Liability (Medium-High Severity)", [
            "Bodily injury / Personenschaden \u2192 LIAB-600 \u2014 Bodily Injury Liability",
            "Third-party property / Sachschaden \u2192 LIAB-610 \u2014 Third-Party Property Damage",
        ]),
        ("Engineering (Medium Severity)", [
            "Machinery breakdown / Maschinenbruch \u2192 ENG-700 \u2014 Machinery Breakdown",
        ]),
        ("Motor & Transport (Medium Severity)", [
            "Motor / Kfz-Schaden \u2192 MOTOR-800 \u2014 Motor Damage",
            "Cargo / Transport \u2192 MARINE-810 \u2014 Marine Cargo",
        ]),
    ]
    for title, items in groups:
        doc.add_heading(title, level=2)
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

# ---------- Required Documents Checklist ----------
def docs_checklist(doc):
    doc.add_heading("Required Documents Checklist", level=1)
    sections = [
        ("Property Damage Claim", [
            "Loss Notification (FNOL) \u2013 Required",
            "Proof of Loss \u2013 Required",
            "Damage Photos \u2013 Required",
            "Repair / Replacement Estimate \u2013 Required",
            "Police / Fire Report (if applicable) \u2013 Required",
        ]),
        ("Liability Claim", [
            "Loss Notification (FNOL) \u2013 Required",
            "Third-Party Demand / Claim Letter \u2013 Required",
            "Incident Report \u2013 Required",
            "Witness Statements \u2013 Required",
            "Loss History (5 years) \u2013 Required",
        ]),
        ("Fast-Track (FT) Claim \u2013 Simple", [
            "Insured Name \u2013 Required",
            "Loss Location / Address \u2013 Required",
            "Date of Loss \u2013 Required",
            "Loss Description \u2013 Required",
            "Estimated Loss Amount (EUR) \u2013 Required",
        ]),
    ]
    for title, items in sections:
        doc.add_heading(title, level=2)
        for i, it in enumerate(items, 1):
            doc.add_paragraph(f"{i}. {it}")

# ---------- Fast-Track Reserve Rules ----------
def reserve_rules(doc):
    doc.add_heading("Fast-Track Reserve Rules", level=1)
    doc.add_heading("Required Fields", level=2)
    doc.add_paragraph("All of the following must be present for a Fast-Track reserve estimate:")
    for f in ["Insured Name", "Estimated Loss Amount", "Loss Location", "Loss Description"]:
        doc.add_paragraph(f, style="List Bullet")
    doc.add_heading("Base Reserve Factors (by severity)", level=2)
    for f in ["Low Severity: 0.60", "Medium Severity: 0.85", "High Severity: 1.10"]:
        doc.add_paragraph(f, style="List Bullet")
    doc.add_heading("Location Factors (by Bundesland)", level=2)
    for f in ["Bayern: 1.00", "Nordrhein-Westfalen: 1.05", "Baden-W\u00fcrttemberg: 1.00",
              "Hessen: 1.05", "Berlin: 1.10", "Hamburg: 1.05"]:
        doc.add_paragraph(f, style="List Bullet")
    doc.add_heading("Reserve Terms", level=2)
    for f in ["Minimum Reserve: \u20ac2,500", "Maximum Fast-Track Reserve: \u20ac50,000 (route larger losses to Standard handling)",
              "Standard Excess (Deductible): \u20ac1,000", "Currency: EUR"]:
        doc.add_paragraph(f, style="List Bullet")
    doc.add_heading("Formula", level=2)
    doc.add_paragraph("Reserve = base_factor \u00d7 estimated_loss \u00d7 location_factor")
    doc.add_paragraph("If the calculated reserve is below the minimum, apply the \u20ac2,500 floor.")
    doc.add_paragraph("Always mark Fast-Track reserves as DRAFT pending claims adjuster review.")

save(r"cr932_uwSubmissionProcessor.file.Region_Mapping.docx_hjD\filedata\Region_Mapping.docx", region)
save(r"cr932_uwSubmissionProcessor.file.Underwriter_Table.docx_Jj2\filedata\Underwriter_Table.docx", adjusters)
save(r"cr932_underwriterAssignmentAgent.file.Underwriter_Table.docx_mG3\filedata\Underwriter_Table.docx", adjusters)
save(r"cr932_uwSubmissionProcessor.file.Business_Classification_Reference.docx_pd-\filedata\Business_Classification_Reference.docx", causes)
save(r"cr932_uwSubmissionProcessor.file.Required_Documents_Checklist.docx_jnN\filedata\Required_Documents_Checklist.docx", docs_checklist)
save(r"cr932_uwSubmissionProcessor.file.ACS_Quote_Rules.docx_vGg\filedata\ACS_Quote_Rules.docx", reserve_rules)
print("DONE")
